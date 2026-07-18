import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docxcompose.composer import Composer
import os
import math
import glob

# ---------------------------------------------------------
# --- Popup Dialog for Excel Sheet Selection ---
# ---------------------------------------------------------
class SheetSelectDialog(tk.Toplevel):
    def __init__(self, parent, sheets):
        super().__init__(parent)
        self.title("Select Excel Sheet")
        self.geometry("350x150")
        self.resizable(False, False)
        self.selected_sheet = None
        
        self.update_idletasks()
        x = parent.winfo_x() + (parent.winfo_width() // 2) - (175)
        y = parent.winfo_y() + (parent.winfo_height() // 2) - (75)
        self.geometry(f"+{x}+{y}")
        
        ttk.Label(self, text="အသုံးပြုလိုသော Sheet ကို ရွေးချယ်ပါ:", font=('Arial', 10, 'bold')).pack(pady=(15, 5))
        self.sheet_var = tk.StringVar(value=sheets[0])
        self.cb = ttk.Combobox(self, textvariable=self.sheet_var, values=sheets, state="readonly", font=('Arial', 10), width=30)
        self.cb.pack(pady=5)
        
        ttk.Button(self, text="OK", command=self.on_ok).pack(pady=10)
        self.transient(parent)
        self.grab_set()
        parent.wait_window(self)
        
    def on_ok(self):
        self.selected_sheet = self.sheet_var.get()
        self.destroy()

class App:
    def __init__(self, root):
        self.root = root
        self.root.title("MKS Professional Mapping Tool")
        self.root.geometry("1000x850")
        
        style = ttk.Style()
        style.theme_use('clam')
        
        self.mapping_rules = {} 
        self.template_path = None
        self.excel_path = None
        self.selected_sheet = None

        # --- Action Bar ---
        action_frame = tk.Frame(root)
        action_frame.pack(side=tk.BOTTOM, fill=tk.X, padx=20, pady=20)
        
        btn_font = ('Arial', 11, 'bold')
        tk.Button(action_frame, text="🚀 GENERATE REPORT", bg="#4CAF50", fg="white", font=btn_font, height=2, command=self.generate_report).pack(side=tk.LEFT, padx=5, expand=True, fill=tk.X)
        tk.Button(action_frame, text="📂 OPEN FOLDER", bg="#2196F3", fg="white", font=btn_font, height=2, command=self.open_folder).pack(side=tk.LEFT, padx=5, expand=True, fill=tk.X)
        tk.Button(action_frame, text="🔄 RESET", bg="#FF9800", fg="white", font=btn_font, height=2, command=self.reset_app).pack(side=tk.LEFT, padx=5, expand=True, fill=tk.X)
        tk.Button(action_frame, text="❌ EXIT", bg="#F44336", fg="white", font=btn_font, height=2, command=self.root.destroy).pack(side=tk.LEFT, padx=5, expand=True, fill=tk.X)

        # --- Main Layout ---
        main_frame = ttk.Frame(root, padding="20")
        main_frame.pack(side=tk.TOP, fill=tk.BOTH, expand=True)

        # File Load Section
        file_frame = ttk.LabelFrame(main_frame, text=" 📂 File Selection ", padding="10")
        file_frame.pack(fill=tk.X, pady=(0, 10))
        ttk.Button(file_frame, text="Load Word Template", command=self.load_template).grid(row=0, column=0, padx=5, pady=5)
        ttk.Button(file_frame, text="Load Excel Data", command=self.load_excel).grid(row=0, column=1, padx=5, pady=5)
        self.lbl_word = ttk.Label(file_frame, text="Word: Not Selected", foreground="gray")
        self.lbl_word.grid(row=0, column=2, padx=10)
        self.lbl_excel = ttk.Label(file_frame, text="Excel: Not Selected", foreground="gray")
        self.lbl_excel.grid(row=0, column=3, padx=10)

        # Mapping Section
        map_frame = ttk.Frame(main_frame)
        map_frame.pack(fill=tk.BOTH, expand=True)

        self.word_all = self.create_box(map_frame, "Word Columns (Available)", 0)
        self.word_sel = self.create_box(map_frame, "Word Selected", 1)
        self.excel_all = self.create_box(map_frame, "Excel Columns (Available)", 2)
        self.excel_sel = self.create_box(map_frame, "Excel Selected", 3)

        # Action Buttons for manual move
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(pady=5)
        ttk.Button(btn_frame, text="Add Rule", command=self.create_rule).pack(side=tk.LEFT, padx=10)

        # Rules List
        rules_pane = ttk.LabelFrame(main_frame, text=" Current Mapping Rules ", padding="10")
        rules_pane.pack(fill=tk.BOTH, expand=True, pady=5)
        self.list_rules = tk.Listbox(rules_pane, height=6, borderwidth=1, relief="solid", font=('Arial', 10))
        self.list_rules.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        ttk.Button(rules_pane, text="🗑️ Remove", command=self.remove_rule).pack(side=tk.LEFT, padx=10)

    # --- UI Helper ---
    def create_box(self, parent, title, col):
        frame = ttk.LabelFrame(parent, text=title, padding="5")
        frame.grid(row=0, column=col, sticky="nsew", padx=5)
        parent.columnconfigure(col, weight=1)
        lb = tk.Listbox(frame, height=10, borderwidth=1, relief="solid", font=('Arial', 10))
        lb.pack(fill=tk.BOTH, expand=True)
        lb.bind('<Double-Button-1>', self.move_double_click)
        return lb

    def move_double_click(self, event):
        lb = event.widget
        selection = lb.curselection()
        if not selection: return
        
        if lb == self.word_all: self.move_item(self.word_all, self.word_sel)
        elif lb == self.word_sel: self.move_item(self.word_sel, self.word_all)
        elif lb == self.excel_all: self.move_item(self.excel_all, self.excel_sel)
        elif lb == self.excel_sel: self.move_item(self.excel_sel, self.excel_all)

    def move_item(self, source, dest):
        selection = source.curselection()
        for i in reversed(selection):
            dest.insert(tk.END, source.get(i))
            source.delete(i)

    # --- Core Logic ---
    def create_rule(self):
        word_header = self.word_sel.get(0) if self.word_sel.size() > 0 else None
        excel_cols = list(self.excel_sel.get(0, tk.END))
        if word_header and excel_cols:
            self.mapping_rules[word_header] = excel_cols
            self.list_rules.insert(tk.END, f"{word_header} <--- {', '.join(excel_cols)}")
            self.word_sel.delete(0, tk.END)
            self.excel_sel.delete(0, tk.END)
        else:
            messagebox.showwarning("Warning", "Word မှ (၁) ခု နှင့် Excel မှ (၁) ခု သို့မဟုတ် ထက်ပို၍ ရွေးချယ်ပေးပါ။")

    def remove_rule(self):
        selection = self.list_rules.curselection()
        if selection:
            idx = selection[0]
            header = self.list_rules.get(idx).split(" <--- ")[0]
            if header in self.mapping_rules: del self.mapping_rules[header]
            self.list_rules.delete(idx)

    def load_excel(self):
        path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if path:
            try:
                xls = pd.ExcelFile(path)
                dialog = SheetSelectDialog(self.root, xls.sheet_names)
                if dialog.selected_sheet:
                    self.excel_path = path
                    self.selected_sheet = dialog.selected_sheet
                    df = pd.read_excel(path, sheet_name=self.selected_sheet)
                    self.excel_all.delete(0, tk.END)
                    self.excel_sel.delete(0, tk.END)
                    for col in df.columns: self.excel_all.insert(tk.END, col)
                    self.lbl_excel.config(text=f"Excel: {os.path.basename(path)}", foreground="green")
            except Exception as e: messagebox.showerror("Error", f"Excel error: {e}")

    def load_template(self):
        path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        if path:
            try:
                self.template_path = path
                doc = Document(path)
                self.word_all.delete(0, tk.END)
                self.word_sel.delete(0, tk.END)
                if doc.tables:
                    for cell in doc.tables[0].rows[0].cells: self.word_all.insert(tk.END, cell.text.strip())
                self.lbl_word.config(text=f"Word: {os.path.basename(path)}", foreground="green")
            except Exception as e: messagebox.showerror("Error", f"Template error: {e}")

    # --- 🛠️ ပြင်ဆင်ပြီး စိတ်ချရသော REPORT GENERATION LOGIC 🛠️ ---
    def generate_report(self):
        if not self.template_path or not self.excel_path:
            messagebox.showwarning("Warning", "Word Template နှင့် Excel Data ဖိုင်များကို အရင် Load လုပ်ပေးပါ။")
            return
        if not self.mapping_rules:
            messagebox.showwarning("Warning", "Mapping Rules များ အနည်းဆုံးတစ်ခု သတ်မှတ်ပေးပါ။")
            return

        # Output Folder ဆောက်ခြင်း
        output_dir = os.path.join(os.getcwd(), "output")
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        output_path = os.path.join(output_dir, "QR Print.docx")

        # ဖိုင်ဖွင့်ထားဆဲ ဖြစ်နေပါက Error အရင်ပြရန်
        if os.path.exists(output_path):
            try:
                # ဖိုင်ကို တခြားသူ ဖွင့်ထားလား စမ်းသပ်ကြည့်ခြင်း
                f = open(output_path, "a")
                f.close()
            except IOError:
                messagebox.showerror("Error", "ထွက်ပေါ်မည့် 'QR Print.docx' ဖိုင်အား Word တွင် ဖွင့်ထားဆဲ ဖြစ်နေပါသည်။\nကျေးဇူးပြု၍ ထို Word ဖိုင်ကို အရင်ပိတ်ပေးပြီးမှ ခလုတ်ကို ထပ်နှိပ်ပါ။")
                return

        try:
            df = pd.read_excel(self.excel_path, sheet_name=self.selected_sheet)
            
            # ပထမဆုံး စာမျက်နှာအတွက် Master Template ဖန်တီးခြင်း
            master_doc = Document(self.template_path)
            composer = Composer(master_doc)
            
            # ဒေတာတစ်ခုချင်းစီကို ကွင်းဆက်မပြတ်ဘဲ စိတ်ချရစွာ လုပ်ဆောင်ရန် Loop ပတ်ခြင်း
            for index, row in df.iterrows():
                # ကူးယူရန် ယာယီ Template တစ်ခုချင်းစီကို ခေါ်ယူခြင်း
