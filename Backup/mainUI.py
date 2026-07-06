import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import pandas as pd
from docx import Document
import os

class App:
    def __init__(self, root):
        self.root = root
        self.root.title("Excel to Word Filling Tool")
        self.root.geometry("800x700")
        
        self.mapping_rules = {} 

        # --- UI Layout ---
        btn_frame = tk.Frame(root)
        btn_frame.pack(pady=10)
        tk.Button(btn_frame, text="Select Word Template", command=self.load_template).pack(side=tk.LEFT, padx=10)
        tk.Button(btn_frame, text="Select Excel File", command=self.load_excel).pack(side=tk.LEFT, padx=10)

        main_frame = tk.Frame(root)
        main_frame.pack(pady=10)

        # Word Headers
        left_frame = tk.LabelFrame(main_frame, text="Word Table Headers")
        left_frame.pack(side=tk.LEFT, padx=10)
        self.word_all = tk.Listbox(left_frame, width=30, height=8)
        self.word_all.pack()
        tk.Button(left_frame, text="ADD", command=lambda: self.move_item(self.word_all, self.word_sel)).pack()
        tk.Button(left_frame, text="REMOVE", command=lambda: self.move_item(self.word_sel, self.word_all)).pack()
        self.word_sel = tk.Listbox(left_frame, width=30, height=2)
        self.word_sel.pack()

        # Excel Columns
        right_frame = tk.LabelFrame(main_frame, text="Excel Sheet Columns")
        right_frame.pack(side=tk.LEFT, padx=10)
        self.excel_all = tk.Listbox(right_frame, width=30, height=8, selectmode=tk.MULTIPLE)
        self.excel_all.pack()
        tk.Button(right_frame, text="ADD", command=lambda: self.move_item(self.excel_all, self.excel_sel)).pack()
        tk.Button(right_frame, text="REMOVE", command=lambda: self.move_item(self.excel_sel, self.excel_all)).pack()
        self.excel_sel = tk.Listbox(right_frame, width=30, height=8, selectmode=tk.MULTIPLE)
        self.excel_sel.pack()

        # Actions
        tk.Button(root, text="CREATE MAPPING RULE", bg="yellow", command=self.create_rule).pack(pady=5)
        self.list_rules = tk.Listbox(root, width=100, height=6)
        self.list_rules.pack(pady=5)
        
        tk.Button(root, text="GENERATE REPORT (Preserve Format)", bg="lightgreen", font=('Arial', 12, 'bold'), command=self.generate_report).pack(pady=20)

    def move_item(self, source, dest):
        selection = source.curselection()
        for i in reversed(selection):
            item = source.get(i)
            dest.insert(tk.END, item)
            source.delete(i)

    def create_rule(self):
        word_header = self.word_sel.get(0)
        excel_cols = list(self.excel_sel.get(0, tk.END))
        if word_header and excel_cols:
            self.mapping_rules[word_header] = excel_cols
            self.list_rules.insert(tk.END, f"{word_header} <--- {', '.join(excel_cols)}")
            self.word_sel.delete(0, tk.END)
            self.excel_sel.delete(0, tk.END)

    def generate_report(self):
        if not hasattr(self, 'template_path'): return
        
        # Open Original Template
        doc = Document(self.template_path)
        df = pd.read_excel(self.excel_path)
        
        table = doc.tables[0] # Assuming first table
        header_row = table.rows[0]
        
        # Get Header indices
        header_map = {cell.text.strip(): i for i, cell in enumerate(header_row.cells)}
        
        # Start filling from Row index 1 (Assuming Row 0 is header)
        for i, row in df.iterrows():
            target_row_idx = i + 1
            
            # If row doesn't exist in template, add new row
            if target_row_idx >= len(table.rows):
                new_row = table.add_row()
            else:
                new_row = table.rows[target_row_idx]
            
            # Fill mapped columns
            for word_header, excel_cols in self.mapping_rules.items():
                if word_header in header_map:
                    col_idx = header_map[word_header]
                    # Join values
                    vals = [str(row[c]) for c in excel_cols if pd.notna(row[c])]
                    new_row.cells[col_idx].text = " ".join(vals)
        
        # Save As New File to protect original template
        if not os.path.exists('output'): os.makedirs('output')
        doc.save("output/Filled_Report.docx")
        messagebox.showinfo("Success", "Format အတိုင်း ဖြည့်ပြီးပါပြီ။ 'output' folder ကိုကြည့်ပါ။")

    def load_excel(self):
        path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        if path:
            self.excel_path = path
            df = pd.read_excel(path)
            self.excel_all.delete(0, tk.END)
            for col in df.columns: self.excel_all.insert(tk.END, col)

    def load_template(self):
        path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        if path:
            self.template_path = path
            doc = Document(path)
            self.word_all.delete(0, tk.END)
            for cell in doc.tables[0].rows[0].cells:
                self.word_all.insert(tk.END, cell.text.strip())

root = tk.Tk()
app = App(root)
root.mainloop()