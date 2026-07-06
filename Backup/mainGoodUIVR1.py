import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import pandas as pd
from docx import Document
import os
import subprocess

class App:
    def __init__(self, root):
        self.root = root
        self.root.title("MKS Professional Mapping Tool")
        self.root.geometry("950x800")
        self.root.minsize(800, 700) # အသေးဆုံးထားနိုင်မည့် Window Size ကို သတ်မှတ်ထားသည်
        
        # UI Styling (Theme)
        style = ttk.Style()
        style.theme_use('clam')
        
        self.mapping_rules = {} 

        # ---------------------------------------------------------
        # --- Bottom Action Bar (အမြဲပေါ်နေစေရန် ပထမဆုံး Pack လုပ်သည်) ---
        # ---------------------------------------------------------
        action_frame = tk.Frame(root)
        action_frame.pack(side=tk.BOTTOM, fill=tk.X, padx=20, pady=20)
        
        btn_font = ('Arial', 11, 'bold')
        
        # ရောင်စုံခလုတ်များ (Colorful Buttons)
        tk.Button(action_frame, text="🚀 GENERATE REPORT", bg="#4CAF50", fg="white", font=btn_font, height=2, command=self.generate_report).pack(side=tk.LEFT, padx=5, expand=True, fill=tk.X)
        tk.Button(action_frame, text="📂 OPEN FOLDER", bg="#2196F3", fg="white", font=btn_font, height=2, command=self.open_folder).pack(side=tk.LEFT, padx=5, expand=True, fill=tk.X)
        tk.Button(action_frame, text="🔄 RESET", bg="#FF9800", fg="white", font=btn_font, height=2, command=self.reset_app).pack(side=tk.LEFT, padx=5, expand=True, fill=tk.X)
        tk.Button(action_frame, text="❌ EXIT", bg="#F44336", fg="white", font=btn_font, height=2, command=self.root.destroy).pack(side=tk.LEFT, padx=5, expand=True, fill=tk.X)

        # ---------------------------------------------------------
        # --- Main Container (အပေါ်ပိုင်း အားလုံးအတွက်) ---
        # ---------------------------------------------------------
        main_frame = ttk.Frame(root, padding="20")
        main_frame.pack(side=tk.TOP, fill=tk.BOTH, expand=True)

        # ၁။ File Selection Section
        file_frame = ttk.LabelFrame(main_frame, text=" 📂 File Selection ", padding="10")
        file_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Button(file_frame, text="Load Word Template", command=self.load_template).grid(row=0, column=0, padx=5, pady=5)
        ttk.Button(file_frame, text="Load Excel Data", command=self.load_excel).grid(row=0, column=1, padx=5, pady=5)
        
        self.lbl_word = ttk.Label(file_frame, text="Word: Not Selected", foreground="gray")
        self.lbl_word.grid(row=0, column=2, padx=10)
        self.lbl_excel = ttk.Label(file_frame, text="Excel: Not Selected", foreground="gray")
        self.lbl_excel.grid(row=0, column=3, padx=10)

        # ၂။ Mapping Section
        map_frame = ttk.Frame(main_frame)
        map_frame.pack(fill=tk.BOTH, expand=True)

        # ဘယ်ဘက် (Word Side)
        left_pane = ttk.LabelFrame(map_frame, text=" Word Columns (Header) ", padding="5")
        left_pane.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 5))
        
        self.word_all = tk.Listbox(left_pane, height=8, borderwidth=1, relief="solid", font=('Arial', 10))
        self.word_all.pack(fill=tk.BOTH, expand=True)
        ttk.Button(left_pane, text="Add Selection ▼", command=lambda: self.move_item(self.word_all, self.word_sel)).pack(fill=tk.X, pady=4)
        
        self.word_sel = tk.Listbox(left_pane, height=2, borderwidth=1, relief="solid", font=('Arial', 10), bg="#e8f4f8")
        self.word_sel.pack(fill=tk.X)
        ttk.Button(left_pane, text="Remove ▲", command=lambda: self.move_item(self.word_sel, self.word_all)).pack(fill=tk.X, pady=(4, 0))

        # ညာဘက် (Excel Side)
        right_pane = ttk.LabelFrame(map_frame, text=" Excel Columns (Data) ", padding="5")
        right_pane.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(5, 0))
        
        self.excel_all = tk.Listbox(right_pane, height=8, selectmode=tk.MULTIPLE, borderwidth=1, relief="solid", font=('Arial', 10))
        self.excel_all.pack(fill=tk.BOTH, expand=True)
        ttk.Button(right_pane, text="Add Selection ▼", command=lambda: self.move_item(self.excel_all, self.excel_sel)).pack(fill=tk.X, pady=4)
        
        self.excel_sel = tk.Listbox(right_pane, height=4, selectmode=tk.MULTIPLE, borderwidth=1, relief="solid", font=('Arial', 10), bg="#e8f4f8")
        self.excel_sel.pack(fill=tk.X)
        ttk.Button(right_pane, text="Remove ▲", command=lambda: self.move_item(self.excel_sel, self.excel_all)).pack(fill=tk.X, pady=(4, 0))

        # ၃။ Rule Creation
        ttk.Button(main_frame, text="➕ Create Mapping Rule", command=self.create_rule).pack(pady=10)

        # ၄။ Rules Display
        rules_pane = ttk.LabelFrame(main_frame, text=" Current Mapping Rules ", padding="10")
        rules_pane.pack(fill=tk.BOTH, expand=True, pady=5)
        
        self.list_rules = tk.Listbox(rules_pane, height=6, borderwidth=1, relief="solid", font=('Arial', 10))
        self.list_rules.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        ttk.Button(rules_pane, text="🗑️ Remove Rule", command=self.remove_rule).pack(side=tk.LEFT, padx=10)


    # ---------------------------------------------------------
    # --- Methods / လုပ်ဆောင်ချက်များ ---
    # ---------------------------------------------------------
    def move_item(self, source, dest):
        selection = source.curselection()
        for i in reversed(selection):
            item = source.get(i)
            dest.insert(tk.END, item)
            source.delete(i)

    def create_rule(self):
        word_header = self.word_sel.get(0) if self.word_sel.size() > 0 else None
        excel_cols = list(self.excel_sel.get(0, tk.END))
        
        if word_header and excel_cols:
            self.mapping_rules[word_header] = excel_cols
            self.list_rules.insert(tk.END, f"{word_header} <--- {', '.join(excel_cols)}")
            self.word_sel.delete(0, tk.END)
            self.excel_sel.delete(0, tk.END)
        else:
            messagebox.showwarning("Warning", "Word မှ (၁) ခု နှင့် Excel မှ (၁) ခု သို့မဟုတ် ထက်ပို၍ ရွေးချယ်ပေးပါ။")

    def remove_rule(self):
        selection = self.list_rules.curselection()
        if selection:
            idx = selection[0]
            item = self.list_rules.get(idx)
            header = item.split(" <--- ")[0]
            if header in self.mapping_rules:
                del self.mapping_rules[header]
            self.list_rules.delete(idx)

    def load_excel(self):
        path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        if path:
            self.excel_path = path
            self.lbl_excel.config(text=f"Excel: {os.path.basename(path)}", foreground="green")
            df = pd.read_excel(path)
            self.excel_all.delete(0, tk.END)
            self.excel_sel.delete(0, tk.END)
            for col in df.columns: self.excel_all.insert(tk.END, col)

    def load_template(self):
        path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        if path:
            self.template_path = path
            self.lbl_word.config(text=f"Word: {os.path.basename(path)}", foreground="green")
            doc = Document(path)
            self.word_all.delete(0, tk.END)
            self.word_sel.delete(0, tk.END)
            # ဇယား၏ ပထမဆုံး row (Header) ကို ဖတ်ယူသည်
            for cell in doc.tables[0].rows[0].cells:
                self.word_all.insert(tk.END, cell.text.strip())

    # --- New Button Functions ---
    def open_folder(self):
        if not os.path.exists('output'):
            os.makedirs('output')
        try:
            os.startfile('output') # Windows အတွက် Folder ဖွင့်ပေးသည်
        except Exception as e:
            messagebox.showerror("Error", f"Folder ဖွင့်၍မရပါ။ \n{e}")

    def reset_app(self):
        # အချက်အလက်များအားလုံးကို အလွတ် (Reset) ပြန်ထားသည်
        self.mapping_rules.clear()
        self.list_rules.delete(0, tk.END)
        self.word_all.delete(0, tk.END)
        self.word_sel.delete(0, tk.END)
        self.excel_all.delete(0, tk.END)
        self.excel_sel.delete(0, tk.END)
        
        if hasattr(self, 'template_path'): del self.template_path
        if hasattr(self, 'excel_path'): del self.excel_path
        
        self.lbl_word.config(text="Word: Not Selected", foreground="gray")
        self.lbl_excel.config(text="Excel: Not Selected", foreground="gray")
        messagebox.showinfo("Reset", "အချက်အလက်အားလုံး အသစ်ပြန်လည်သတ်မှတ်ပြီးပါပြီ။")

    # --- Document Generator Logic (Strict Format အတိုင်း) ---
    def generate_report(self):
        if not hasattr(self, 'template_path') or not hasattr(self, 'excel_path'):
            messagebox.showerror("Error", "Word Template နှင့် Excel Data ဖိုင်များ ရွေးပေးပါ။")
            return
        if not self.mapping_rules:
            messagebox.showerror("Error", "Mapping Rule တစ်ခုခု ဖန်တီးပေးပါ။")
            return

        try:
            df = pd.read_excel(self.excel_path)
            master_doc = Document(self.template_path)
            new_doc = Document()
            
            table = master_doc.tables[0]
            rows_per_page = len(table.rows) - 1

            before_table = [p for p in master_doc.paragraphs if p._element.index < table._element.index]
            after_table = [p for p in master_doc.paragraphs if p._element.index > table._element.index]

            for i in range(0, len(df), rows_per_page):
                batch = df.iloc[i : i + rows_per_page]
                if i > 0: new_doc.add_page_break()

                # 1. Letterhead / Heading ကို ကူးယူသည်
                for p in before_table:
                    new_doc.add_paragraph(p.text, style=p.style)
                
                # 2. Table Format ကို ကူးယူသည်
                new_table = new_doc.add_table(rows=1, cols=len(table.columns[0].cells))
                new_table.style = table.style
                for j, cell in enumerate(table.rows[0].cells):
                    new_table.rows[0].cells[j].text = cell.text

                # 3. Data များဖြည့်သွင်းသည်
                for _, row in batch.iterrows():
                    new_row = new_table.add_row()
                    for col_idx, cell in enumerate(table.rows[0].cells):
                        header = cell.text.strip()
                        if header in self.mapping_rules:
                            excel_cols = self.mapping_rules[header]
                            vals = [str(row[c]) for c in excel_cols if pd.notna(row[c])]
                            new_row.cells[col_idx].text = " ".join(vals)
                        else:
                            new_row.cells[col_idx].text = ""

                # 4. Signatures / Footer ကို ကူးယူသည်
                for p in after_table:
                    new_doc.add_paragraph(p.text, style=p.style)

            if not os.path.exists('output'): os.makedirs('output')
            new_doc.save("output/Final_Report.docx")
            messagebox.showinfo("Success", "Format မပျက်ဘဲ ဖိုင်အောင်မြင်စွာ ထွက်ရှိပါပြီ။\n'OPEN FOLDER' ကို နှိပ်၍ ဖွင့်ကြည့်နိုင်ပါသည်။")
            
        except Exception as e:
            messagebox.showerror("Error", f"လုပ်ဆောင်ချက် မအောင်မြင်ပါ။\n{e}")

root = tk.Tk()
app = App(root)
root.mainloop()