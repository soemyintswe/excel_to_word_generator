import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
import os
import math
import glob

# ---------------------------------------------------------
# --- Popup Dialog for Excel Sheet Selection ---
# ---------------------------------------------------------
class SheetSelectDialog(tk.Toplevel):
    def __init__(self, parent, sheets):
        super().__init__(parent)
        self.title("Select Excel Sheet")
        self.geometry("350x150")
        self.resizable(False, False)
        self.selected_sheet = None
        
        # Center the dialog window
        self.update_idletasks()
        x = parent.winfo_x() + (parent.winfo_width() // 2) - (350 // 2)
        y = parent.winfo_y() + (parent.winfo_height() // 2) - (150 // 2)
        self.geometry(f"+{x}+{y}")
        
        ttk.Label(self, text="အသုံးပြုလိုသော Sheet ကို ရွေးချယ်ပါ:", font=('Arial', 10, 'bold')).pack(pady=(15, 5))
        
        self.sheet_var = tk.StringVar(value=sheets[0])
        self.cb = ttk.Combobox(self, textvariable=self.sheet_var, values=sheets, state="readonly", font=('Arial', 10), width=30)
        self.cb.pack(pady=5)
        
        ttk.Button(self, text="OK", command=self.on_ok).pack(pady=10)
        
        self.transient(parent)
        self.grab_set()
        parent.wait_window(self)
        
    def on_ok(self):
        self.selected_sheet = self.sheet_var.get()
        self.destroy()


class App:
    def __init__(self, root):
        self.root = root
        self.root.title("MKS Professional Mapping Tool")
        self.root.geometry("950x800")
        self.root.minsize(800, 700) 
        
        style = ttk.Style()
        style.theme_use('clam')
        
        self.mapping_rules = {} 

        # ---------------------------------------------------------
        # --- Bottom Action Bar ---
        # ---------------------------------------------------------
        action_frame = tk.Frame(root)
        action_frame.pack(side=tk.BOTTOM, fill=tk.X, padx=20, pady=20)
        
        btn_font = ('Arial', 11, 'bold')
        
        tk.Button(action_frame, text="🚀 GENERATE REPORT", bg="#4CAF50", fg="white", font=btn_font, height=2, command=self.generate_report).pack(side=tk.LEFT, padx=5, expand=True, fill=tk.X)
        tk.Button(action_frame, text="📂 OPEN FOLDER", bg="#2196F3", fg="white", font=btn_font, height=2, command=self.open_folder).pack(side=tk.LEFT, padx=5, expand=True, fill=tk.X)
        tk.Button(action_frame, text="🔄 RESET", bg="#FF9800", fg="white", font=btn_font, height=2, command=self.reset_app).pack(side=tk.LEFT, padx=5, expand=True, fill=tk.X)
        tk.Button(action_frame, text="❌ EXIT", bg="#F44336", fg="white", font=btn_font, height=2, command=self.root.destroy).pack(side=tk.LEFT, padx=5, expand=True, fill=tk.X)

        # ---------------------------------------------------------
        # --- Main Container ---
        # ---------------------------------------------------------
        main_frame = ttk.Frame(root, padding="20")
        main_frame.pack(side=tk.TOP, fill=tk.BOTH, expand=True)

        file_frame = ttk.LabelFrame(main_frame, text=" 📂 File Selection ", padding="10")
        file_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Button(file_frame, text="Load Word Template", command=self.load_template).grid(row=0, column=0, padx=5, pady=5)
        ttk.Button(file_frame, text="Load Excel Data", command=self.load_excel).grid(row=0, column=1, padx=5, pady=5)
        
        self.lbl_word = ttk.Label(file_frame, text="Word: Not Selected", foreground="gray")
        self.lbl_word.grid(row=0, column=2, padx=10)
        self.lbl_excel = ttk.Label(file_frame, text="Excel: Not Selected", foreground="gray")
        self.lbl_excel.grid(row=0, column=3, padx=10)

        map_frame = ttk.Frame(main_frame)
        map_frame.pack(fill=tk.BOTH, expand=True)

        left_pane = ttk.LabelFrame(map_frame, text=" Word Columns (Header) ", padding="5")
        left_pane.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 5))
        
        self.word_all = tk.Listbox(left_pane, height=8, borderwidth=1, relief="solid", font=('Arial', 10))
        self.word_all.pack(fill=tk.BOTH, expand=True)
        ttk.Button(left_pane, text="Add Selection ▼", command=lambda: self.move_item(self.word_all, self.word_sel)).pack(fill=tk.X, pady=4)
        
        self.word_sel = tk.Listbox(left_pane, height=2, borderwidth=1, relief="solid", font=('Arial', 10), bg="#e8f4f8")
        self.word_sel.pack(fill=tk.X)
        ttk.Button(left_pane, text="Remove ▲", command=lambda: self.move_item(self.word_sel, self.word_all)).pack(fill=tk.X, pady=(4, 0))

        right_pane = ttk.LabelFrame(map_frame, text=" Excel Columns (Data) ", padding="5")
        right_pane.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(5, 0))
        
        self.excel_all = tk.Listbox(right_pane, height=8, selectmode=tk.MULTIPLE, borderwidth=1, relief="solid", font=('Arial', 10))
        self.excel_all.pack(fill=tk.BOTH, expand=True)
        ttk.Button(right_pane, text="Add Selection ▼", command=lambda: self.move_item(self.excel_all, self.excel_sel)).pack(fill=tk.X, pady=4)
        
        self.excel_sel = tk.Listbox(right_pane, height=4, selectmode=tk.MULTIPLE, borderwidth=1, relief="solid", font=('Arial', 10), bg="#e8f4f8")
        self.excel_sel.pack(fill=tk.X)
        ttk.Button(right_pane, text="Remove ▲", command=lambda: self.move_item(self.excel_sel, self.excel_all)).pack(fill=tk.X, pady=(4, 0))

        ttk.Button(main_frame, text="➕ Create Mapping Rule", command=self.create_rule).pack(pady=10)

        rules_pane = ttk.LabelFrame(main_frame, text=" Current Mapping Rules ", padding="10")
        rules_pane.pack(fill=tk.BOTH, expand=True, pady=5)
        
        self.list_rules = tk.Listbox(rules_pane, height=6, borderwidth=1, relief="solid", font=('Arial', 10))
        self.list_rules.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        ttk.Button(rules_pane, text="🗑️ Remove Rule", command=self.remove_rule).pack(side=tk.LEFT, padx=10)

    # ---------------------------------------------------------
    # --- Methods ---
    # ---------------------------------------------------------
    def move_item(self, source, dest):
        selection = source.curselection()
        for i in reversed(selection):
            item = source.get(i)
            dest.insert(tk.END, item)
            source.delete(i)

    def create_rule(self):
        word_header = self.word_sel.get(0) if self.word_sel.size() > 0 else None
        excel_cols = list(self.excel_sel.get(0, tk.END))
        
        if word_header and excel_cols:
            self.mapping_rules[word_header] = excel_cols
            self.list_rules.insert(tk.END, f"{word_header} <--- {', '.join(excel_cols)}")
            self.word_sel.delete(0, tk.END)
            self.excel_sel.delete(0, tk.END)
        else:
            messagebox.showwarning("Warning", "Word မှ (၁) ခု နှင့် Excel မှ (၁) ခု သို့မဟုတ် ထက်ပို၍ ရွေးချယ်ပေးပါ။")

    def remove_rule(self):
        selection = self.list_rules.curselection()
        if selection:
            idx = selection[0]
            item = self.list_rules.get(idx)
            header = item.split(" <--- ")[0]
            if header in self.mapping_rules:
                del self.mapping_rules[header]
            self.list_rules.delete(idx)

    # ဤနေရာတွင် Sheet ရွေးချယ်နိုင်သော စနစ်ကို အသစ်ထည့်သွင်းထားပါသည်
    def load_excel(self):
        path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if path:
            try:
                xls = pd.ExcelFile(path)
                sheets = xls.sheet_names
                
                # Sheet ရွေးရန် Popup ခေါ်ခြင်း
                dialog = SheetSelectDialog(self.root, sheets)
                selected_sheet = dialog.selected_sheet
                
                # OK နှိပ်ပြီး Sheet ရွေးခဲ့မှသာ ဆက်လက်လုပ်ဆောင်မည်
                if selected_sheet:
                    self.excel_path = path
                    self.selected_sheet = selected_sheet
                    self.lbl_excel.config(text=f"Excel: {os.path.basename(path)} [{selected_sheet}]", foreground="green")
                    
                    df = pd.read_excel(path, sheet_name=selected_sheet)
                    self.excel_all.delete(0, tk.END)
                    self.excel_sel.delete(0, tk.END)
                    
                    for col in df.columns: 
                        self.excel_all.insert(tk.END, col)
            except Exception as e:
                messagebox.showerror("Error", f"Excel ဖိုင်ဖွင့်ရာတွင် အမှားဖြစ်နေပါသည်။\n{e}")

    def load_template(self):
        path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        if path:
            self.template_path = path
            self.lbl_word.config(text=f"Word: {os.path.basename(path)}", foreground="green")
            doc = Document(path)
            self.word_all.delete(0, tk.END)
            self.word_sel.delete(0, tk.END)
            if doc.tables:
                for cell in doc.tables[0].rows[0].cells:
                    self.word_all.insert(tk.END, cell.text.strip())
            else:
                messagebox.showerror("Error", "ဤ Word ဖိုင်ထဲတွင် ဇယား (Table) မတွေ့ပါ။")

    def open_folder(self):
        if not os.path.exists('output'):
            os.makedirs('output')
        try:
            os.startfile('output') 
        except Exception as e:
            messagebox.showerror("Error", f"Folder ဖွင့်၍မရပါ။ \n{e}")

    def reset_app(self):
        self.mapping_rules.clear()
        self.list_rules.delete(0, tk.END)
        self.word_all.delete(0, tk.END)
        self.word_sel.delete(0, tk.END)
        self.excel_all.delete(0, tk.END)
        self.excel_sel.delete(0, tk.END)
        
        if hasattr(self, 'template_path'): del self.template_path
        if hasattr(self, 'excel_path'): del self.excel_path
        if hasattr(self, 'selected_sheet'): del self.selected_sheet
        
        self.lbl_word.config(text="Word: Not Selected", foreground="gray")
        self.lbl_excel.config(text="Excel: Not Selected", foreground="gray")
        messagebox.showinfo("Reset", "အချက်အလက်အားလုံး အသစ်ပြန်လည်သတ်မှတ်ပြီးပါပြီ။")

    # ---------------------------------------------------------
    # --- Document Generator Logic (In-Place Edit + Formatting) ---
    # ---------------------------------------------------------
    def generate_report(self):
        if not hasattr(self, 'template_path') or not hasattr(self, 'excel_path'):
            messagebox.showerror("Error", "Word Template နှင့် Excel Data ဖိုင်များ ရွေးပေးပါ။")
            return
        if not self.mapping_rules:
            messagebox.showerror("Error", "Mapping Rule တစ်ခုခု ဖန်တီးပေးပါ။")
            return

        try:
            # ၁။ Output Folder အသစ်ဆောက်ခြင်း (သို့) ဖိုင်ဟောင်းများ ရှင်းလင်းခြင်း
            if not os.path.exists('output'):
                os.makedirs('output')
            else:
                for f in glob.glob("output/*.docx"):
                    try:
                        os.remove(f)
                    except Exception:
                        pass # ဖွင့်ထားသော ဖိုင်များရှိလျှင် ကျော်သွားမည်

            # ရွေးချယ်ထားသော Sheet ကိုသာ ဖတ်ယူရန်
            df = pd.read_excel(self.excel_path, sheet_name=getattr(self, 'selected_sheet', 0))
            
            temp_doc = Document(self.template_path)
            table_template = temp_doc.tables[0]
            max_data_rows = len(table_template.rows) - 1
            
            if max_data_rows <= 0:
                messagebox.showerror("Error", "ဇယားထဲတွင် အချက်အလက်ဖြည့်ရန် အလွတ် Row များ မရှိပါ။")
                return

            total_batches = math.ceil(len(df) / max_data_rows)
            
            for batch_idx in range(total_batches):
                start_idx = batch_idx * max_data_rows
                end_idx = start_idx + max_data_rows
                batch_df = df.iloc[start_idx:end_idx]
                
                doc = Document(self.template_path)
                table = doc.tables[0]
                
                for i, (_, row) in enumerate(batch_df.iterrows()):
                    row_idx_in_word = i + 1 
                    
                    for col_idx, cell in enumerate(table.rows[0].cells):
                        header = cell.text.strip()
                        if header in self.mapping_rules:
                            excel_cols = self.mapping_rules[header]
                            vals = [str(row[c]) for c in excel_cols if pd.notna(row[c])]
                            final_val = " ".join(vals)
                            
                            target_cell = table.rows[row_idx_in_word].cells[col_idx]
                            target_cell.text = "" # မူလစာသားများကို ရှင်းလင်းမည်
                            
                            # Formatting နှင့် Font များ ထည့်သွင်းခြင်း
                            p = target_cell.paragraphs[0]
                            run = p.add_run(final_val)
                            
                            # Pyidaungsu Numbers သတ်မှတ်ခြင်း
                            run.font.name = "Pyidaungsu Numbers"
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Pyidaungsu Numbers')
                            run._element.rPr.rFonts.set(qn('w:cs'), 'Pyidaungsu Numbers')
                            
                            # Spacing သတ်မှတ်ခြင်း
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.space_after = Pt(0)
                            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

                output_filename = f"output/Report_Page_{batch_idx + 1}.docx"
                doc.save(output_filename)

            messagebox.showinfo("Success", f"Format နှင့် Font များ အတိအကျဖြင့် ဖိုင် ({total_batches}) ဖိုင် အောင်မြင်စွာ ထွက်ရှိပါပြီ။\n'OPEN FOLDER' ကို နှိပ်၍ ဖွင့်ကြည့်နိုင်ပါသည်။")
            
        except Exception as e:
            messagebox.showerror("Error", f"လုပ်ဆောင်ချက် မအောင်မြင်ပါ။\n{e}")

root = tk.Tk()
app = App(root)
root.mainloop()